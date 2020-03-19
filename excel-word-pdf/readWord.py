import docx
import os

os.chdir('/home/daniel/Downloads')

d = docx.Document('demo.docx')

print(d.paragraphs[0].text)

p = d.paragraphs[1]

# A run is a segment of styled text (e.g. bold, italic)
print(p.runs[1].text)  # bold

# We can change styles and text:
p.runs[3].underline = True
p.runs[3].text = 'italic and underline'

d.save('demo2.docx')
