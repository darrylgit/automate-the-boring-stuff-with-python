import docx
import os

os.chdir('/home/daniel/Downloads')

d = docx.Document('demo4.docx')
# d.add_paragraph('Hello this is a paragraph.')
# d.add_paragraph('This is another paragraph.')

p = d.paragraphs[0]
p.add_run('This is a new run.')
p.runs[1].bold = True

d.save('demo4.docx')
